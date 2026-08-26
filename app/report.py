"""
Builds the metareview report docx — same section structure as Sega's existing Metareview
Report template, rendering the output of narrative.draft_narrative() rather than a placeholder.
Uses python-docx (pure Python) so the whole app runs on one runtime.
"""
from io import BytesIO
from typing import Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_report_docx(game_title: str, methodology_note: str, narrative: dict,
                       review_count: int, average_score: Optional[float],
                       disclosures: list, platform_breakdown: Optional[dict] = None) -> bytes:
    """
    narrative: the dict returned by narrative.draft_narrative() — executive_summary (str),
    press_reactions_synthesis (list[str]), category_callouts (list of {category, label,
    label_caveat, synthesis, quotes: [{text, outlet}]}), recommendations (list of {heading, text}),
    fan_reaction (list[str], empty when this run had no fan/user notes supplied).
    disclosures: always render these — this tool's output is a first draft, not a final report.
    platform_breakdown: output of matrix.compute_platform_averages(), or None/empty — rendered
    as extra lines under Review Averages only when 2+ distinct platforms are represented (a
    single-platform run just gets the existing blended Review count / Average score line, same
    as the real Persona 3: Reload reference report, which shows no platform breakdown at all).
    """
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading(f"{game_title} — Metareview", level=0).alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(methodology_note)
    doc.add_paragraph(
        "A critique is defined as a statement of positivity or negativity towards an aspect "
        "of the game. A reviewer who makes both a positive and a negative statement about the "
        "same aspect receives a neutral rating. A reviewer who only mentions a feature without "
        "opinion is not counted. Categories are only reported when at least 25% of the sample "
        "commented on them. This metareview only covers press/critic reviews, which differs "
        "from user/fan feedback."
    )

    doc.add_heading("Executive Summary", level=1)
    for para in narrative.get("executive_summary", "").split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.add_heading("Review Averages", level=1)
    p = doc.add_paragraph()
    score_text = f"{average_score:.2f}" if average_score is not None else "N/A (no review in this batch stated a numeric score)"
    p.add_run(f"Review count: {review_count}    ·    Average score: {score_text}").bold = True

    if platform_breakdown and len(platform_breakdown) >= 2:
        # Matches the real Puyo Puyo Tetris reference report's per-storefront/platform Review
        # Averages breakdown — the real P3R report, by contrast, shows no breakdown at all, which
        # is exactly what happens here too when a run has fewer than 2 distinct platforms tagged.
        for platform, stats in platform_breakdown.items():
            plat_score_text = (
                f"{stats['average_score']:.2f}" if stats["average_score"] is not None else "N/A"
            )
            doc.add_paragraph(
                f"{platform}: {stats['review_count']} review(s), average score {plat_score_text}",
                style="List Bullet",
            )

    doc.add_heading("Reactions", level=1)
    doc.add_heading("Press Reactions", level=2)
    for para in narrative.get("press_reactions_synthesis", []):
        doc.add_paragraph(para)

    doc.add_heading("The Good and the Bad", level=2)
    for co in narrative.get("category_callouts", []):
        h = doc.add_paragraph()
        h.add_run(f"{co['category']} [{co['label']}]").bold = True
        if co.get("label_caveat"):
            cav = doc.add_paragraph()
            cav.add_run(f"*Note: {co['label_caveat']}").italic = True
        if co.get("synthesis"):
            doc.add_paragraph(co["synthesis"])
        for q in co.get("quotes", []):
            qp = doc.add_paragraph()
            qp.add_run(f"“{q['text']}” — {q['outlet']}").italic = True

    fan_reaction = narrative.get("fan_reaction", [])
    if fan_reaction:
        # Omitted entirely when empty (no producer-supplied fan/user notes this run) — same
        # "don't show it empty" principle as the recommendations section below. Matches the
        # placement and content of the real Puyo Puyo Tetris report's "Impression of Fan
        # Reaction" section, which comes right after the category call-outs.
        doc.add_heading("Impression of Fan Reaction", level=2)
        doc.add_paragraph(
            "Unlike the press-review sentiment above, this section is not derived from the "
            "same threshold-scored methodology — it reflects a producer-supplied sample of "
            "fan/user commentary (forums, storefront user reviews, survey responses, etc.) and "
            "should be treated as impressionistic, not statistically representative."
        )
        for para in fan_reaction:
            doc.add_paragraph(para)

    recommendations = narrative.get("recommendations", [])
    if recommendations:
        # Omitted entirely (not shown with an empty heading) when there's nothing under it —
        # same precedent as the reference Puyo Puyo Tetris report explicitly omitting its
        # "Review" sub-section for a title it didn't apply to, rather than showing it empty.
        doc.add_heading("Review and Recommendations", level=1)
        doc.add_heading("PD Recommendations for Next Title", level=2)
        for rec in recommendations:
            rh = doc.add_paragraph()
            rh.add_run(rec["heading"]).bold = True
            doc.add_paragraph(rec["text"])

    doc.add_heading("Metareview Matrix", level=1)
    doc.add_paragraph("See the attached matrix workbook (Summary tab) for the full category "
                       "breakdown, formulas, and opinion graph.")

    doc.add_heading("Notes on this AI-assisted run", level=1)
    for d in disclosures:
        doc.add_paragraph(d, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()