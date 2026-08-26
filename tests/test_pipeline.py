"""
Exercises the plumbing (extraction, matrix building, report building, weighted-score math)
with synthetic/mocked data — no live network calls and no Anthropic API key required. This
validates that a real run's non-LLM machinery is correct; the classify.py Claude calls
themselves need a live ANTHROPIC_API_KEY to test for real (see README).

Run: pytest tests/test_pipeline.py -v
"""
import sys
from pathlib import Path
from types import SimpleNamespace

sys.path.insert(0, str(Path(__file__).parent.parent))

from app import metacritic
from app.classify import (
    _normalize_classifications, _normalize_review_language, _normalize_stated_score,
    classify_review, classify_reviews_for_categories, detect_emergent_categories,
)
from app.extraction import extract_score_from_text, extract_text_from_html
from app.matrix import (
    build_matrix_workbook, compute_platform_averages, compute_weighted_scores,
    suggest_sentiment_label,
)
from app.metacritic import (
    build_sources_from_metacritic, build_sources_from_metacritic_all_platforms, fetch_full_text,
    fetch_game_platforms, fetch_review_list, fetch_review_list_all_platforms,
    parse_game_url_or_slug,
)
from app.narrative import _build_narrative_tool, draft_narrative
from app.report import build_report_docx
from app.rubric import is_on_curated_list, load_bundled_curated_list


class _FakeMessages:
    """Stands in for anthropic.Anthropic().messages so classify.py's Claude calls are testable
    without a live API key — returns a single forced tool_use block with whatever input dict
    the test wants to simulate."""
    def __init__(self, block_input, tool_name):
        self._block_input, self._tool_name = block_input, tool_name

    def create(self, **kwargs):
        block = SimpleNamespace(type="tool_use", name=self._tool_name, input=self._block_input)
        return SimpleNamespace(content=[block])


class _FakeClient:
    def __init__(self, block_input, tool_name="submit_classification"):
        self.messages = _FakeMessages(block_input, tool_name)


class _QueuedMessages:
    """Like _FakeMessages, but returns a different queued response on each successive call and
    records the kwargs each call was made with — for testing narrative.draft_narrative()'s
    retry-on-empty-result behavior."""
    def __init__(self, responses, tool_name="submit_narrative"):
        self._responses = list(responses)
        self._tool_name = tool_name
        self.calls = []

    def create(self, **kwargs):
        self.calls.append(kwargs)
        block_input = self._responses.pop(0)
        block = SimpleNamespace(type="tool_use", name=self._tool_name, input=block_input)
        return SimpleNamespace(content=[block])


class _QueuedClient:
    def __init__(self, responses, tool_name="submit_narrative"):
        self.messages = _QueuedMessages(responses, tool_name)


def test_extract_text_from_html_strips_boilerplate():
    html = """
    <html><body>
      <nav>Home | Reviews | Contact</nav>
      <article>
        <h1>Test Game Review</h1>
        <p>The combat in this game is fast and satisfying, a clear step up from the original.</p>
        <p>The story, however, drags in the middle third and never quite recovers.</p>
      </article>
      <footer>Copyright 2026</footer>
    </body></html>
    """
    text = extract_text_from_html(html)
    assert text is not None
    assert "combat" in text.lower()
    assert "Home | Reviews | Contact" not in text


def test_curated_list_lookup():
    outlets, snapshot_date = load_bundled_curated_list()
    assert len(outlets) > 100
    assert snapshot_date
    assert is_on_curated_list("IGN", outlets)
    assert is_on_curated_list(" ign ", outlets)  # case/whitespace insensitive
    assert not is_on_curated_list("Definitely Not A Real Outlet", outlets)


def test_normalize_classifications_handles_expected_shape():
    raw = [
        {"category": "Combat", "value": 1, "quote": "great combat"},
        {"category": "Story", "value": -1, "quote": "weak story"},
    ]
    assert _normalize_classifications(raw) == raw


def test_normalize_classifications_drops_bare_strings_without_crashing():
    # This exact shape — a list mixing a bare category-name string in with real objects — is
    # what crashed a live 50-review run with "TypeError: string indices must be integers, not
    # 'str'" in streamlit_app.py before this fix: the model deviated from the forced tool
    # schema on one review, and the old code trusted the shape unconditionally.
    raw = [
        {"category": "Combat", "value": 1, "quote": "great combat"},
        "Story",  # malformed: no value/quote attached
        {"category": "Visuals", "value": 0, "quote": "mixed bag"},
    ]
    assert _normalize_classifications(raw) == [
        {"category": "Combat", "value": 1, "quote": "great combat"},
        {"category": "Visuals", "value": 0, "quote": "mixed bag"},
    ]


def test_normalize_classifications_handles_dict_shaped_response():
    # Some models occasionally return a category->value mapping instead of an array of objects.
    raw = {"Combat": {"value": 1, "quote": "great combat"}, "Story": -1}
    by_category = {c["category"]: c for c in _normalize_classifications(raw)}
    assert by_category["Combat"] == {"category": "Combat", "value": 1, "quote": "great combat"}
    assert by_category["Story"] == {"category": "Story", "value": -1, "quote": ""}


def test_classify_review_end_to_end_survives_malformed_model_output():
    # Full path through classify_review() with a fake client returning the malformed shape that
    # crashed the real run, then the exact downstream loop from streamlit_app.py, to confirm it
    # no longer raises.
    fake_input = {
        "classifications": [
            {"category": "Combat", "value": 1, "quote": "great combat"},
            "Story",
        ],
        "candidate_emergent_topics": ["Pacing", 42, "Recasting"],
    }
    result = classify_review(_FakeClient(fake_input), "some review text", ["Combat", "Story"])

    data = {}
    for c in result["classifications"]:  # the exact line that used to crash
        data.setdefault(c["category"], {})["review_key"] = (c["value"], c["quote"])
    assert data == {"Combat": {"review_key": (1, "great combat")}}
    assert result["candidate_emergent_topics"] == ["Pacing", "Recasting"]


def test_extract_score_from_text_various_formats():
    # The numerator/denominator notation itself is language-agnostic — the same regex has to
    # catch a French/Spanish/German outlet's score without any per-language keyword.
    cases = [
        ("Overall this is a great game. 8.5/10 in my book.", 85.0),
        ("Note : 8/10 — un excellent jeu.", 80.0),          # French
        ("Nota: 8,5/10, un juego excelente.", 85.0),         # Spanish, comma decimal
        ("Wertung: 90/100", 90.0),                            # German
        ("I'd give this a solid 4/5 stars.", 80.0),
        ("Overall score: 90%", 90.0),  # keyword must precede the number — see tier-3 note below
        ("A stellar 90% score from us.", None),  # keyword AFTER the number doesn't count (below)
        ("Posted on 8/10/2024, this game earns high marks throughout.", None),  # US date
        ("Released 08/10/24 — no explicit score given, just prose.", None),     # short date
        ("This game has no numeric score at all, purely narrative review text.", None),
        (None, None),
        ("", None),
    ]
    for text, expected in cases:
        assert extract_score_from_text(text) == expected, f"text={text!r}"


def test_extract_score_from_text_real_pdf_extracted_snippets():
    # Every case here is a verbatim (or near-verbatim) snippet from pypdf's actual text-layer
    # extraction of 15 real Persona 3 Reload review PDFs — not synthetic examples. Each comment
    # is a genuine failure mode this heuristic had to be tuned against, not a hypothetical:
    cases = [
        # A real score, correctly labeled, on one line — the easy case.
        ("• Story: 7/10\n• Combat: 9/10\n• Presentation: 7/10\nScore: 8/10", 80.0),
        # A brand name sitting between the "Verdict" label and the score, with no separating
        # space before the following word either (pypdf lost the whitespace on both sides).
        ("Our VerdictPCGamesN 8/10Whether you're a new player experiencing", 80.0),
        # A decimal number that pypdf split into two pieces with a stray space at the decimal
        # point, immediately butted up against "out of" with no space at all.
        ("-Poor balance of vocal music tracks\n7 .3out of 10Persona 3 Reload Review", 73.0),
        # A bare integer directly after "Score" (no slash, no %) — assumed out of 10.
        ("Score 10The Final Word\nPersona 3 Reload is a flawless", 100.0),
        # An age rating right next to the word "Rating" — must NOT be read as a review score.
        ("RatingPEGI 16 (https://www.psu.com/rating/pegi-16/)\nScore 10", 100.0),
        # An unrelated discount percentage — must NOT be read as a review score.
        ("H28.7\nPersona 3 ReloadSviluppato da Atlus. Prodotto da SEGA.71% di sconto", None),
        # An unrelated ad percentage — must NOT be read as a review score.
        ("TRENDING Get Surfshark VPN 86% OFF \nWindows 11 \nID@Xbox", None),
        # A date that happens to contain "/10" — must NOT be read as a review score.
        ("more coffee than AlanWake.\n23/12/25\n21/10/25\n18/12/24\n02/02/24Persona 3 Reload", None),
        # A URL ending in a platform code, immediately followed on the next line by an unrelated
        # page-footer digit — the two must NOT bridge into a bogus fraction across the linebreak.
        ("persona-3-reload-review-ps5/\n5 of 14 8/12/2026, 2:05 PM", None),
        # A score rendered as icon-font glyphs (private-use-area codepoints) has no digits to
        # find at all — correctly unextractable without OCR/icon interpretation, not a bug.
        ("Windows Central Verdict  \nIn 2006, Atlus released", None),
    ]
    for text, expected in cases:
        got = extract_score_from_text(text)
        assert got == expected, f"expected={expected!r}, got={got!r}, text={text!r}"


def test_normalize_review_language_and_stated_score_helpers():
    assert _normalize_review_language("French") == "French"
    assert _normalize_review_language("  ") == "Unknown"
    assert _normalize_review_language(None) == "Unknown"
    assert _normalize_review_language(123) == "Unknown"

    assert _normalize_stated_score(85) == 85.0
    assert _normalize_stated_score(85.5) == 85.5
    assert _normalize_stated_score(None) is None
    assert _normalize_stated_score(150) is None  # out of range — don't trust it
    assert _normalize_stated_score(-5) is None
    assert _normalize_stated_score(True) is None  # bool is an int subclass in Python — reject it


def test_classify_review_returns_language_and_stated_score():
    fake_input = {
        "classifications": [{"category": "Combat", "value": 1, "quote": "Combat great."}],
        "candidate_emergent_topics": [],
        "review_language": "French",
        "stated_score": 85,
        "stated_score_raw": "8.5/10",
    }
    result = classify_review(_FakeClient(fake_input), "Un excellent jeu.", ["Combat"])
    assert result["review_language"] == "French"
    assert result["stated_score"] == 85.0
    assert result["stated_score_raw"] == "8.5/10"


def test_classify_review_handles_malformed_language_and_score_fields():
    # Same defensive stance as the classifications fix — a malformed field on these two new
    # attributes shouldn't be able to crash a run either.
    fake_input = {
        "classifications": [],
        "candidate_emergent_topics": [],
        "review_language": None,        # model omitted it
        "stated_score": "8.5/10",       # model gave the raw string instead of a normalized number
        "stated_score_raw": 42,          # wrong type
    }
    result = classify_review(_FakeClient(fake_input), "text", ["Combat"])
    assert result["review_language"] == "Unknown"
    assert result["stated_score"] is None
    assert result["stated_score_raw"] is None


def test_normalize_classifications_filters_out_categories_not_in_valid_list():
    # Defense-in-depth for the enum constraint in _build_classify_tool: even if a backend didn't
    # enforce the schema's enum as strictly as the direct Anthropic API does, a paraphrased
    # category name should still be dropped here rather than silently creating an untracked key
    # that never reaches the matrix or report.
    raw = [
        {"category": "Combat", "value": 1, "quote": "great combat"},
        {"category": "Missing Content", "value": -1, "quote": "paraphrased, not the exact name"},
    ]
    valid = ["Combat", "Missing Content from Prior Versions (FES/Portable/The Answer)"]
    assert _normalize_classifications(raw, valid_categories=valid) == [
        {"category": "Combat", "value": 1, "quote": "great combat"},
    ]
    # Without the valid_categories argument, existing behavior (no filtering) is unchanged.
    assert len(_normalize_classifications(raw)) == 2


def test_classify_review_drops_paraphrased_category_not_in_given_list():
    # Reproduces the real bug: confirmed against a live 35-review run where every long, compound
    # custom category name (e.g. "Missing Content from Prior Versions (FES/Portable/The
    # Answer)") came back with ZERO classifications, while short default categories worked fine
    # — because the model paraphrased the long names back slightly differently than given, and
    # the mismatched key silently never reached the matrix. classify_review must not let a
    # paraphrased category name survive into its returned classifications.
    given_categories = ["Combat", "Missing Content from Prior Versions (FES/Portable/The Answer)"]
    fake_input = {
        "classifications": [
            {"category": "Combat", "value": 1, "quote": "great combat"},
            {"category": "Missing Content from Prior Versions", "value": -1,  # paraphrased
             "quote": "no answer epilogue this time"},
        ],
        "candidate_emergent_topics": [],
    }
    result = classify_review(_FakeClient(fake_input), "review text", given_categories)
    assert result["classifications"] == [{"category": "Combat", "value": 1, "quote": "great combat"}]


def test_build_classify_tool_constrains_category_to_given_enum():
    from app.classify import _build_classify_tool
    categories = ["Combat", "Story"]
    tool = _build_classify_tool(categories)
    category_schema = tool["input_schema"]["properties"]["classifications"]["items"]["properties"]["category"]
    assert category_schema["enum"] == categories


def test_detect_emergent_categories_survives_malformed_cluster_entries():
    fake_input = {"clusters": [
        {"label": "Pacing", "mention_count": 8},
        "Recasting",  # malformed: no mention_count attached
        {"label": "Combat", "mention_count": 1},
    ]}
    client = _FakeClient(fake_input, tool_name="submit_clusters")
    emergent = detect_emergent_categories(client, ["pacing"] * 8, num_reviews=10, threshold=0.25)
    assert emergent == ["Pacing"]  # 8/10=0.8 clears the threshold; Combat's 1/10 doesn't


def test_classify_reviews_for_categories_merges_results_and_reports_progress():
    # Regression for a real bug: emergent categories were being named (via
    # detect_emergent_categories) and added to the matrix, but no review was ever actually
    # classified against them — every stat came back flatly 0, indistinguishable from genuine
    # unanimous silence on the topic. This is the follow-up pass that fixes that; this test
    # confirms it actually produces real per-review scored data, merged by category.
    sources = [
        SimpleNamespace(key="a", outlet="Outlet A", text="review a text"),
        SimpleNamespace(key="b", outlet="Outlet B", text="review b text"),
    ]
    responses = [
        {"classifications": [{"category": "Theurgy", "value": 1, "quote": "great new mechanic"}],
         "candidate_emergent_topics": [], "review_language": "English",
         "stated_score": None, "stated_score_raw": None},
        {"classifications": [{"category": "Theurgy", "value": -1, "quote": "too strong"}],
         "candidate_emergent_topics": [], "review_language": "English",
         "stated_score": None, "stated_score_raw": None},
    ]
    client = _QueuedClient(responses, tool_name="submit_classification")

    progress_calls = []
    data, failed = classify_reviews_for_categories(
        client, sources, ["Theurgy"], model="test-model",
        progress_cb=lambda done, total: progress_calls.append((done, total)),
    )

    assert failed == []
    assert data["Theurgy"]["a"] == (1, "great new mechanic")
    assert data["Theurgy"]["b"] == (-1, "too strong")
    assert progress_calls == [(1, 2), (2, 2)]


def test_classify_reviews_for_categories_isolates_failures_per_review():
    sources = [
        SimpleNamespace(key="a", outlet="Outlet A", text="review a text"),
        SimpleNamespace(key="b", outlet="Outlet B", text="review b text"),
    ]

    class _FlakyMessages:
        def __init__(self):
            self.calls = 0

        def create(self, **kwargs):
            self.calls += 1
            if self.calls == 1:
                raise RuntimeError("simulated API error")
            block = SimpleNamespace(type="tool_use", name="submit_classification", input={
                "classifications": [{"category": "Theurgy", "value": 1, "quote": "loved it"}],
                "candidate_emergent_topics": [], "review_language": "English",
                "stated_score": None, "stated_score_raw": None,
            })
            return SimpleNamespace(content=[block])

    client = SimpleNamespace(messages=_FlakyMessages())
    data, failed = classify_reviews_for_categories(client, sources, ["Theurgy"], model="test-model")

    assert failed == [("Outlet A", "simulated API error")]
    assert "a" not in data.get("Theurgy", {})
    assert data["Theurgy"]["b"] == (1, "loved it")


def _sample_reviews_and_data():
    reviews = [
        {"key": "a", "outlet": "Outlet A", "score": 90, "date": "1/1/24"},
        {"key": "b", "outlet": "Outlet B", "score": 80, "date": "1/2/24"},
        {"key": "c", "outlet": "Outlet C", "score": 70, "date": "1/3/24"},
        {"key": "d", "outlet": "Outlet D", "score": 95, "date": "1/4/24"},
    ]
    categories = ["Combat", "Story", "Rarely Mentioned"]
    data = {
        "Combat": {"a": (1, "great combat"), "b": (1, "fun combat"), "c": (-1, "clunky combat")},
        "Story": {"a": (-1, "weak story"), "b": (0, "story is both moving and repetitive"), "d": (1, "great story")},
        "Rarely Mentioned": {"a": (1, "one lone mention")},
    }
    return reviews, categories, data


def test_compute_weighted_scores_matches_expected_formula():
    reviews, categories, data = _sample_reviews_and_data()
    scores = compute_weighted_scores(reviews, categories, data, threshold=0.25)

    # Combat: 2 positive, 1 negative, out of 4 reviews -> (2-1)/4 = 0.25
    assert scores["Combat"]["weighted"] == 0.25
    assert scores["Combat"]["meets_threshold"] is True

    # Story: 1 pos, 1 neutral, 1 neg, out of 4 -> (1-1)/4 = 0.0, mention_rate 3/4 = 0.75 >= 0.25
    assert scores["Story"]["weighted"] == 0.0
    assert scores["Story"]["meets_threshold"] is True

    # Rarely Mentioned: 1 mention out of 4 = 0.25 -> exactly at threshold, should meet it
    assert scores["Rarely Mentioned"]["mention_rate"] == 0.25
    assert scores["Rarely Mentioned"]["meets_threshold"] is True


def test_suggest_sentiment_label_matches_reference_reports():
    # Real (positive, mixed, negative) counts pulled from the Puyo Puyo Tetris metareview
    # matrix, checked against the bracketed label the human-written report actually used for
    # that category. This is a real regression check, not synthetic data.
    cases = [
        (10, 0, 0, "Universally Praised"),      # Breadth of Content
        (6, 0, 0, "Universally Praised"),        # Switch Features
        (9, 0, 1, "Widely Praised"),             # Combining Puyo Puyo and Tetris
        (7, 1, 0, "Widely Praised"),              # Aesthetic
        (4, 1, 0, "Widely Praised"),              # Music
        (4, 2, 0, "Generally Praised"),           # Big Bang
        (8, 4, 7, "Extremely Controversial"),     # Adventure Mode Story
        (3, 4, 2, "Extremely Controversial"),     # Voice Acting
    ]
    for pos, mixed, neg, expected in cases:
        assert suggest_sentiment_label(pos, mixed, neg) == expected, \
            f"({pos},{mixed},{neg}) expected {expected!r}, got {suggest_sentiment_label(pos, mixed, neg)!r}"

    # Sanity checks on the symmetric (negative-leaning) side, which the reference reports
    # happened not to exercise with a clean example:
    assert suggest_sentiment_label(0, 0, 5) == "Universally Panned"
    assert suggest_sentiment_label(1, 0, 2) in ("Controversial, Leaning Negative",
                                                 "Somewhat Controversial, Leaning Negative")
    assert suggest_sentiment_label(0, 0, 0) == "Not enough data"


def test_build_matrix_workbook_has_formulas_not_hardcoded_values():
    reviews, categories, data = _sample_reviews_and_data()
    wb = build_matrix_workbook("Test Game", reviews, categories, data)
    ws = wb["Summary"]
    assert ws["A1"].value == "Test Game — Metareview"
    assert str(ws["A7"].value).startswith("=")  # weighted-score formula, not a literal number
    assert str(ws["E2"].value).startswith("=COUNTA")


def test_build_matrix_workbook_review_count_and_weighted_formula_use_outlet_row_not_score_row():
    # Regression for a real bug: Review Count and the weighted-score formula's denominator were
    # both counting non-blank cells in the SCORE row (row 5) instead of the always-populated
    # outlet-name row (row 2). Confirmed against a live 35-review run where every review's score
    # was blank: "Review Count" showed 0 and every category's weighted score showed #DIV/0!,
    # even though 35 real, included reviews were plainly present in the sheet.
    reviews, categories, data = _sample_reviews_and_data()
    for r in reviews:
        r["score"] = None  # simulate a batch where no review has a stated score
    wb = build_matrix_workbook("Test Game", reviews, categories, data)
    ws = wb["Summary"]
    # Row 2 = outlet names (always populated); row 5 = score (can be legitimately all-blank).
    assert "2:" in str(ws["E2"].value) and "5:" not in str(ws["E2"].value)
    assert str(ws["E2"].value).startswith("=COUNTA")
    assert "$2:" in str(ws["A7"].value) and "$5:" not in str(ws["A7"].value)
    # Average Score must degrade gracefully (IFERROR -> "N/A") rather than surfacing a raw
    # #DIV/0! to a producer opening the spreadsheet.
    assert str(ws["E3"].value).startswith("=IFERROR")


def test_compute_platform_averages_groups_scores_and_skips_untagged_reviews():
    reviews = [
        {"key": "a", "score": 90, "platform": "PlayStation 5"},
        {"key": "b", "score": 80, "platform": "PlayStation 5"},
        {"key": "c", "score": 70, "platform": "PC"},
        {"key": "d", "score": None, "platform": "PC"},  # no stated score — still counted
        {"key": "e", "score": 60, "platform": None},  # untagged — excluded entirely
        {"key": "f", "score": 50, "platform": "  "},  # blank-ish — also excluded
    ]
    result = compute_platform_averages(reviews)

    # PS5 has more reviews (2 vs PC's 2 too — tie broken alphabetically: PC < PlayStation 5).
    assert list(result.keys()) == ["PC", "PlayStation 5"]
    assert result["PlayStation 5"] == {"review_count": 2, "average_score": 85.0}
    # PC has a review with no stated score — still counted in review_count, but that review
    # doesn't drag the average toward 0 (it's excluded from the average, not treated as a 0).
    assert result["PC"] == {"review_count": 2, "average_score": 70.0}


def test_compute_platform_averages_reports_na_not_zero_when_platform_has_no_scores_at_all():
    reviews = [
        {"key": "a", "score": None, "platform": "Nintendo Switch"},
        {"key": "b", "score": None, "platform": "Nintendo Switch"},
    ]
    result = compute_platform_averages(reviews)
    assert result["Nintendo Switch"] == {"review_count": 2, "average_score": None}


def test_draft_narrative_retries_when_callouts_and_recommendations_both_empty():
    # Regression for a real bug: on a real 35-review run with 10+ well-supported categories, the
    # model filled in a rich Executive Summary and Press Reactions synthesis (the first two
    # schema fields) and then returned completely empty category_callouts and recommendations —
    # likely running low on output budget by the time it reached the later array fields.
    category_scores = {
        "Combat": {"meets_threshold": True, "weighted": 0.6, "positive": 20, "mixed": 2,
                   "negative": 1, "mention_rate": 0.9, "suggested_label": "Widely Praised"},
    }
    empty_response = {
        "executive_summary": "Well received overall.",
        "press_reactions_synthesis": ["Reviewers praised the combat."],
        "category_callouts": [],
        "recommendations": [],
    }
    populated_response = {
        "executive_summary": "Well received overall.",
        "press_reactions_synthesis": ["Reviewers praised the combat."],
        "category_callouts": [{"category": "Combat", "label": "Widely Praised",
                                "label_caveat": "", "synthesis": "Praised.",
                                "quotes": [{"text": "great combat", "outlet": "Outlet A"}]}],
        "recommendations": [{"heading": "Keep it up", "text": "Maintain combat quality."}],
    }
    client = _QueuedClient([empty_response, populated_response])
    result = draft_narrative(client, "Test Game", "PS5", 85.0, 35, category_scores, {}, [],
                              model="test-model", max_tokens=4000)

    assert result == populated_response  # took the retry's result, not the empty first response
    assert len(client.messages.calls) == 2  # exactly one retry, not a retry loop
    assert client.messages.calls[0]["max_tokens"] == 4000
    assert client.messages.calls[1]["max_tokens"] == 8000  # doubled on retry


def test_draft_narrative_does_not_retry_when_result_already_populated():
    category_scores = {
        "Combat": {"meets_threshold": True, "weighted": 0.6, "positive": 20, "mixed": 2,
                   "negative": 1, "mention_rate": 0.9, "suggested_label": "Widely Praised"},
    }
    populated_response = {
        "executive_summary": "Well received overall.",
        "press_reactions_synthesis": ["Reviewers praised the combat."],
        "category_callouts": [{"category": "Combat", "label": "Widely Praised",
                                "label_caveat": "", "synthesis": "Praised.",
                                "quotes": [{"text": "great combat", "outlet": "Outlet A"}]}],
        "recommendations": [{"heading": "Keep it up", "text": "Maintain combat quality."}],
    }
    client = _QueuedClient([populated_response])
    result = draft_narrative(client, "Test Game", "PS5", 85.0, 35, category_scores, {}, [],
                              model="test-model")
    assert result == populated_response
    assert len(client.messages.calls) == 1  # no wasted retry call


def test_build_narrative_tool_omits_recommendations_property_when_disabled():
    tool_with = _build_narrative_tool(True)
    tool_without = _build_narrative_tool(False)

    assert "recommendations" in tool_with["input_schema"]["properties"]
    assert "recommendations" in tool_with["input_schema"]["required"]

    assert "recommendations" not in tool_without["input_schema"]["properties"]
    assert "recommendations" not in tool_without["input_schema"]["required"]


def test_draft_narrative_forces_empty_recommendations_and_does_not_retry_spuriously():
    # The model is given a schema with no "recommendations" property at all when the flag is
    # off, so it never returns one — that absence must NOT be mistaken for the empty-result
    # failure mode the retry logic above exists to catch (only category_callouts should gate
    # the retry now).
    category_scores = {
        "Combat": {"meets_threshold": True, "weighted": 0.6, "positive": 20, "mixed": 2,
                   "negative": 1, "mention_rate": 0.9, "suggested_label": "Widely Praised"},
    }
    populated_response_no_recommendations = {
        "executive_summary": "Well received overall.",
        "press_reactions_synthesis": ["Reviewers praised the combat."],
        "category_callouts": [{"category": "Combat", "label": "Widely Praised",
                                "label_caveat": "", "synthesis": "Praised.",
                                "quotes": [{"text": "great combat", "outlet": "Outlet A"}]}],
        # no "recommendations" key at all — this is what the model actually returns when the
        # schema doesn't ask for one.
    }
    client = _QueuedClient([populated_response_no_recommendations])
    result = draft_narrative(client, "Test Game", "PS5", 85.0, 35, category_scores, {}, [],
                              model="test-model", include_recommendations=False)

    assert len(client.messages.calls) == 1  # no spurious retry triggered by the missing field
    assert result["recommendations"] == []  # enforced regardless of what the model returned
    assert "recommendations" not in client.messages.calls[0]["tools"][0]["input_schema"]["properties"]


def test_build_narrative_tool_omits_fan_reaction_property_by_default():
    tool_default = _build_narrative_tool(True)  # include_fan_reaction defaults to False
    tool_with = _build_narrative_tool(True, include_fan_reaction=True)

    assert "fan_reaction" not in tool_default["input_schema"]["properties"]
    assert "fan_reaction" in tool_with["input_schema"]["properties"]
    assert "fan_reaction" in tool_with["input_schema"]["required"]


def test_draft_narrative_includes_fan_reaction_only_when_notes_supplied():
    # Whether the "Impression of Fan Reaction" section gets drafted is derived from whether
    # fan_reaction_notes was actually given — not a separate flag — so there's no way to end up
    # with a fan-reaction section drafted from nothing.
    category_scores = {
        "Combat": {"meets_threshold": True, "weighted": 0.6, "positive": 20, "mixed": 2,
                   "negative": 1, "mention_rate": 0.9, "suggested_label": "Widely Praised"},
    }
    populated_response_with_fan_reaction = {
        "executive_summary": "Well received overall.",
        "press_reactions_synthesis": ["Reviewers praised the combat."],
        "category_callouts": [{"category": "Combat", "label": "Widely Praised",
                                "label_caveat": "", "synthesis": "Praised.",
                                "quotes": [{"text": "great combat", "outlet": "Outlet A"}]}],
        "recommendations": [],
        "fan_reaction": ["Fans on the forums echoed the same praise for combat."],
    }
    client = _QueuedClient([populated_response_with_fan_reaction])
    result = draft_narrative(client, "Test Game", "PS5", 85.0, 35, category_scores, {}, [],
                              model="test-model", include_recommendations=False,
                              fan_reaction_notes="  Some forum posts praised combat.  ")

    assert result["fan_reaction"] == ["Fans on the forums echoed the same praise for combat."]
    schema_props = client.messages.calls[0]["tools"][0]["input_schema"]["properties"]
    assert "fan_reaction" in schema_props
    assert "Some forum posts praised combat." in client.messages.calls[0]["messages"][0]["content"]


def test_draft_narrative_forces_empty_fan_reaction_when_no_notes_given():
    category_scores = {
        "Combat": {"meets_threshold": True, "weighted": 0.6, "positive": 20, "mixed": 2,
                   "negative": 1, "mention_rate": 0.9, "suggested_label": "Widely Praised"},
    }
    populated_response = {
        "executive_summary": "Well received overall.",
        "press_reactions_synthesis": ["Reviewers praised the combat."],
        "category_callouts": [{"category": "Combat", "label": "Widely Praised",
                                "label_caveat": "", "synthesis": "Praised.",
                                "quotes": [{"text": "great combat", "outlet": "Outlet A"}]}],
        "recommendations": [],
    }
    for notes in (None, "", "   "):
        client = _QueuedClient([dict(populated_response)])
        result = draft_narrative(client, "Test Game", "PS5", 85.0, 35, category_scores, {}, [],
                                  model="test-model", include_recommendations=False,
                                  fan_reaction_notes=notes)
        assert result["fan_reaction"] == []
        assert "fan_reaction" not in client.messages.calls[0]["tools"][0]["input_schema"]["properties"]


def test_build_report_docx_handles_none_average_score():
    # Regression for a real bug: when no review in a batch has a stated score, average_score
    # used to render as a misleading "0.00" (reading as a universally-panned game) instead of
    # a clear "N/A".
    reviews, categories, data = _sample_reviews_and_data()
    narrative = {
        "executive_summary": "Summary.", "press_reactions_synthesis": [],
        "category_callouts": [], "recommendations": [],
    }
    docx_bytes = build_report_docx(
        game_title="Test Game", methodology_note="note", narrative=narrative,
        review_count=4, average_score=None, disclosures=["test"],
    )
    assert len(docx_bytes) > 1000
    assert docx_bytes[:2] == b"PK"

    from docx import Document
    from io import BytesIO
    doc = Document(BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "N/A" in full_text
    assert "0.00" not in full_text


def test_build_report_docx_produces_nonempty_bytes():
    reviews, categories, data = _sample_reviews_and_data()
    scores = compute_weighted_scores(reviews, categories, data)
    narrative = {
        "executive_summary": "Combat was well received.\n\nStory reception was mixed.",
        "press_reactions_synthesis": [
            "Reviewers consistently praised the combat system as a highlight.",
        ],
        "category_callouts": [
            {
                "category": "Combat",
                "label": scores["Combat"]["suggested_label"],
                "label_caveat": "",
                "synthesis": "Widely praised as fast and satisfying.",
                "quotes": [{"text": "great combat", "outlet": "Outlet A"}],
            },
            {
                "category": "Story",
                "label": scores["Story"]["suggested_label"],
                "label_caveat": "Only a handful of reviews commented directly on the story.",
                "synthesis": "Reception was mixed, with some praise and some criticism.",
                "quotes": [{"text": "weak story", "outlet": "Outlet A"}],
            },
        ],
        "recommendations": [
            {"heading": "Keep Combat's Pacing", "text": "Maintain the current combat pacing in the next title."},
        ],
    }
    docx_bytes = build_report_docx(
        game_title="Test Game",
        methodology_note="4 reviews used for this test.",
        narrative=narrative,
        review_count=4,
        average_score=83.75,
        disclosures=["This is a test run, not a real metareview."],
    )
    assert len(docx_bytes) > 1000  # a real docx, not an empty/broken file
    assert docx_bytes[:2] == b"PK"  # docx is a zip archive


def test_build_report_docx_renders_platform_breakdown_only_with_two_or_more_platforms():
    narrative = {
        "executive_summary": "Summary.", "press_reactions_synthesis": [],
        "category_callouts": [], "recommendations": [],
    }
    from docx import Document
    from io import BytesIO

    # A single platform: the report should stay exactly as the current blended-average form —
    # no per-platform lines, matching the real P3R reference report (which shows no breakdown).
    single_platform = {"PlayStation 5": {"review_count": 4, "average_score": 83.75}}
    docx_bytes = build_report_docx(
        game_title="Test Game", methodology_note="note", narrative=narrative,
        review_count=4, average_score=83.75, disclosures=["test"],
        platform_breakdown=single_platform,
    )
    full_text = "\n".join(p.text for p in Document(BytesIO(docx_bytes)).paragraphs)
    assert "PlayStation 5" not in full_text

    # Two-plus platforms: matches the Puyo Puyo Tetris reference report's per-platform breakdown.
    multi_platform = {
        "PlayStation 5": {"review_count": 3, "average_score": 85.0},
        "PC": {"review_count": 1, "average_score": None},
    }
    docx_bytes = build_report_docx(
        game_title="Test Game", methodology_note="note", narrative=narrative,
        review_count=4, average_score=83.75, disclosures=["test"],
        platform_breakdown=multi_platform,
    )
    full_text = "\n".join(p.text for p in Document(BytesIO(docx_bytes)).paragraphs)
    assert "PlayStation 5: 3 review(s), average score 85.00" in full_text
    assert "PC: 1 review(s), average score N/A" in full_text


def test_build_report_docx_omits_fan_reaction_section_when_empty_but_renders_when_present():
    from docx import Document
    from io import BytesIO

    base_narrative = {
        "executive_summary": "Summary.", "press_reactions_synthesis": [],
        "category_callouts": [], "recommendations": [],
    }

    docx_bytes = build_report_docx(
        game_title="Test Game", methodology_note="note",
        narrative={**base_narrative, "fan_reaction": []},
        review_count=4, average_score=83.75, disclosures=["test"],
    )
    full_text = "\n".join(p.text for p in Document(BytesIO(docx_bytes)).paragraphs)
    assert "Impression of Fan Reaction" not in full_text

    docx_bytes = build_report_docx(
        game_title="Test Game", methodology_note="note",
        narrative={**base_narrative, "fan_reaction": ["Fans loved the combat too."]},
        review_count=4, average_score=83.75, disclosures=["test"],
    )
    full_text = "\n".join(p.text for p in Document(BytesIO(docx_bytes)).paragraphs)
    assert "Impression of Fan Reaction" in full_text
    assert "Fans loved the combat too." in full_text
    assert "impressionistic" in full_text.lower()


# ---------------------------------------------------------------------------------------------
# Metacritic bulk-sourcing (app/metacritic.py) — all HTTP is mocked, no live network calls.
# ---------------------------------------------------------------------------------------------

def test_parse_game_url_or_slug_handles_full_url_with_platform():
    slug, platform = parse_game_url_or_slug(
        "https://www.metacritic.com/game/persona-3-reload/critic-reviews/?platform=xbox-series-x"
    )
    assert slug == "persona-3-reload"
    assert platform == "xbox-series-x"


def test_parse_game_url_or_slug_handles_url_without_platform():
    slug, platform = parse_game_url_or_slug(
        "https://www.metacritic.com/game/persona-3-reload/critic-reviews/"
    )
    assert slug == "persona-3-reload"
    assert platform is None


def test_parse_game_url_or_slug_handles_bare_slug():
    slug, platform = parse_game_url_or_slug("persona-3-reload")
    assert slug == "persona-3-reload"
    assert platform is None


class _FakeReviewListResponse:
    def __init__(self, json_data):
        self._json_data = json_data

    def raise_for_status(self):
        pass

    def json(self):
        return self._json_data


def test_fetch_review_list_pages_until_empty(monkeypatch):
    # Regression guard for the real API behavior found via view-source: it returns (at most)
    # page_size items per call regardless of what's requested, so pagination must continue by
    # walking `offset` forward by however many items actually came back, stopping only when a
    # page comes back genuinely empty — not by trusting a single "total" figure up front.
    pages = [
        {"data": {"items": [{"url": "https://a"}, {"url": "https://b"}]}},
        {"data": {"items": [{"url": "https://c"}]}},
        {"data": {"items": []}},
    ]
    offsets_requested = []

    class _FakeSession:
        def __init__(self):
            self.headers = {}

        def get(self, url, params=None, timeout=None):
            offsets_requested.append(params["offset"])
            return _FakeReviewListResponse(pages.pop(0))

    monkeypatch.setattr(metacritic.requests, "Session", _FakeSession)

    items = fetch_review_list("some-slug", pause_s=0)

    assert [i["url"] for i in items] == ["https://a", "https://b", "https://c"]
    assert offsets_requested == [0, 2, 3]


def test_fetch_review_list_filters_platform_via_url_path_not_query_param(monkeypatch):
    # Regression for a real bug: platform filtering was originally implemented as a
    # `?platform=<slug>` query parameter, which was never actually confirmed against a live
    # response — it turns out that param has no effect at all and the API always returns the
    # lead platform's reviews regardless. Confirmed directly against Metacritic's real API that
    # the filter is instead a `/platform/<slug>/` URL PATH segment (see this function's
    # docstring and app/metacritic.py's module docstring for how this was verified). This test
    # guards against silently regressing back to the query-param version.
    requested = []

    class _FakeSession:
        def __init__(self):
            self.headers = {}

        def get(self, url, params=None, timeout=None):
            requested.append((url, params))
            if len(requested) == 1:
                return _FakeReviewListResponse({"data": {"items": [{"url": "https://xbox-review"}]}})
            return _FakeReviewListResponse({"data": {"items": []}})

    monkeypatch.setattr(metacritic.requests, "Session", _FakeSession)

    items = fetch_review_list("persona-3-reload", platform="xbox-series-x", pause_s=0)

    assert [i["url"] for i in items] == ["https://xbox-review"]
    url, params = requested[0]
    assert url == (
        "https://backend.metacritic.com/reviews/metacritic/critic/games/"
        "persona-3-reload/platform/xbox-series-x/web"
    )
    assert "platform" not in params  # must NOT be sent as a query param — confirmed it's a no-op


def test_fetch_full_text_prefers_requests_tier(monkeypatch):
    monkeypatch.setattr(metacritic, "_try_requests",
                         lambda url, timeout: ("full article text", None))
    monkeypatch.setattr(metacritic, "_try_playwright",
                         lambda url, timeout: (None, "should never be reached"))

    text, method, error = fetch_full_text("https://example.com/review")

    assert text == "full article text"
    assert method == "requests"
    assert error is None


def test_fetch_full_text_falls_back_to_playwright_when_requests_empty(monkeypatch):
    # The real-world case this exists for: an outlet whose article body is rendered
    # client-side, so a bare requests.get() has nothing for trafilatura to extract.
    monkeypatch.setattr(metacritic, "_try_requests", lambda url, timeout: (None, "403 Forbidden"))
    monkeypatch.setattr(metacritic, "_try_playwright",
                         lambda url, timeout: ("js-rendered article text", None))

    text, method, error = fetch_full_text("https://example.com/review")

    assert text == "js-rendered article text"
    assert method == "playwright"
    assert error is None


def test_fetch_full_text_reports_combined_error_when_both_tiers_fail(monkeypatch):
    # Neither "blocked" nor "no Playwright installed" should look like a silent success —
    # both reasons need to reach the producer so they know whether to retry, use the upload
    # tab instead, or install Playwright in this deployment.
    monkeypatch.setattr(metacritic, "_try_requests", lambda url, timeout: (None, "403 Forbidden"))
    monkeypatch.setattr(metacritic, "_try_playwright",
                         lambda url, timeout: (None, "Playwright not installed"))

    text, method, error = fetch_full_text("https://example.com/review")

    assert text is None
    assert method == "none"
    assert "403 Forbidden" in error
    assert "Playwright not installed" in error


class _FakePlaywrightBrowser:
    def __init__(self, html):
        self._html = html

    def new_page(self, user_agent=None):
        return SimpleNamespace(
            goto=lambda url, wait_until=None, timeout=None: None,
            wait_for_timeout=lambda ms: None,
            content=lambda: self._html,
        )

    def close(self):
        pass


class _FakeChromium:
    """launch_effects: a list consumed one per .launch() call — an Exception instance to raise,
    or anything else to return a fake browser instead."""
    def __init__(self, launch_effects, html="<html><body><p>Full article text long enough to "
                                             "survive trafilatura's extraction heuristics here."
                                             "</p></body></html>"):
        self._effects = list(launch_effects)
        self._html = html
        self.launch_calls = 0

    def launch(self, headless=True):
        self.launch_calls += 1
        effect = self._effects.pop(0)
        if isinstance(effect, Exception):
            raise effect
        return _FakePlaywrightBrowser(self._html)


class _FakeSyncPlaywrightCM:
    def __init__(self, chromium):
        self.chromium = chromium

    def __enter__(self):
        return self

    def __exit__(self, *exc_info):
        return False


def test_try_playwright_lazily_installs_chromium_then_retries_launch(monkeypatch):
    # Regression/feature test for Streamlit Community Cloud (and similar hosts with no
    # build-hook to run a separate `playwright install chromium` step): `pip install playwright`
    # gets you the Python package but not the browser binary, so the first real launch fails
    # with Playwright's specific "Executable doesn't exist" error. _try_playwright should catch
    # exactly that, install the browser once, and retry the launch — not treat it as a generic
    # failure.
    monkeypatch.setattr(metacritic, "_playwright_install_state",
                         {"attempted": False, "error": None})
    chromium = _FakeChromium([
        RuntimeError("BrowserType.launch: Executable doesn't exist at /root/.cache/ms-playwright/..."),
        "ok",  # second .launch() call succeeds
    ])
    import playwright.sync_api
    monkeypatch.setattr(playwright.sync_api, "sync_playwright",
                         lambda: _FakeSyncPlaywrightCM(chromium))
    monkeypatch.setattr(metacritic, "_ensure_playwright_chromium_installed", lambda: None)

    text, error = metacritic._try_playwright("https://example.com/review", timeout=20)

    assert error is None
    assert text is not None and "Full article text" in text
    assert chromium.launch_calls == 2  # failed once, retried once after the lazy install


def test_try_playwright_reports_install_failure_without_retrying_launch(monkeypatch):
    monkeypatch.setattr(metacritic, "_playwright_install_state",
                         {"attempted": False, "error": None})
    chromium = _FakeChromium([
        RuntimeError("BrowserType.launch: Executable doesn't exist at /root/.cache/ms-playwright/..."),
    ])
    import playwright.sync_api
    monkeypatch.setattr(playwright.sync_api, "sync_playwright",
                         lambda: _FakeSyncPlaywrightCM(chromium))
    monkeypatch.setattr(metacritic, "_ensure_playwright_chromium_installed",
                         lambda: "playwright install chromium failed: no space left on device")

    text, error = metacritic._try_playwright("https://example.com/review", timeout=20)

    assert text is None
    assert error == "playwright install chromium failed: no space left on device"
    assert chromium.launch_calls == 1  # no wasted retry once the install itself failed


def test_try_playwright_does_not_intercept_unrelated_launch_errors(monkeypatch):
    # Only the specific "browser not installed" failure should trigger the install-and-retry
    # path — any other launch error (e.g. a real crash, an OOM kill) should surface normally,
    # not be silently swallowed while the code waits on an install that was never the problem.
    monkeypatch.setattr(metacritic, "_playwright_install_state",
                         {"attempted": False, "error": None})
    chromium = _FakeChromium([RuntimeError("Target page, context or browser has been closed")])
    import playwright.sync_api
    monkeypatch.setattr(playwright.sync_api, "sync_playwright",
                         lambda: _FakeSyncPlaywrightCM(chromium))
    install_called = []
    monkeypatch.setattr(metacritic, "_ensure_playwright_chromium_installed",
                         lambda: install_called.append(True))

    text, error = metacritic._try_playwright("https://example.com/review", timeout=20)

    assert text is None
    assert "Target page, context or browser has been closed" in error
    assert install_called == []  # never attempted — this wasn't a missing-browser error
    assert chromium.launch_calls == 1


def test_ensure_playwright_chromium_installed_caches_result_across_calls(monkeypatch):
    monkeypatch.setattr(metacritic, "_playwright_install_state",
                         {"attempted": False, "error": None})
    run_calls = []

    def _fake_run(cmd, check=True, capture_output=True, text=True, timeout=300):
        run_calls.append(cmd)
        return SimpleNamespace(returncode=0)

    monkeypatch.setattr(metacritic.subprocess, "run", _fake_run)

    first = metacritic._ensure_playwright_chromium_installed()
    second = metacritic._ensure_playwright_chromium_installed()

    assert first is None
    assert second is None
    assert len(run_calls) == 1  # second call reused the cached success, no repeat subprocess


def test_ensure_playwright_chromium_installed_reports_and_caches_failure(monkeypatch):
    monkeypatch.setattr(metacritic, "_playwright_install_state",
                         {"attempted": False, "error": None})
    run_calls = []

    def _fake_run(cmd, check=True, capture_output=True, text=True, timeout=300):
        run_calls.append(cmd)
        raise metacritic.subprocess.CalledProcessError(
            returncode=1, cmd=cmd, output="", stderr="E: No space left on device\n"
        )

    monkeypatch.setattr(metacritic.subprocess, "run", _fake_run)

    first = metacritic._ensure_playwright_chromium_installed()
    second = metacritic._ensure_playwright_chromium_installed()

    assert first is not None and "No space left on device" in first
    assert second == first  # cached, not re-run
    assert len(run_calls) == 1


def test_build_sources_from_metacritic_maps_fields_and_reports_progress(monkeypatch):
    fake_items = [
        {"url": "https://a.example/review", "publicationName": "Outlet A", "score": 80,
         "date": "2024-01-30"},
        {"url": None, "publicationName": "Outlet B", "score": None, "date": "2024-02-01"},
    ]
    monkeypatch.setattr(metacritic, "fetch_review_list",
                         lambda slug, platform=None: fake_items)
    monkeypatch.setattr(metacritic, "fetch_full_text",
                         lambda url, timeout=20: ("Full article text.", "requests", None))

    progress_calls = []
    sources = build_sources_from_metacritic(
        "some-slug",
        progress_cb=lambda done, total, outlet: progress_calls.append((done, total, outlet)),
    )

    assert len(sources) == 2
    assert sources[0].outlet == "Outlet A"
    assert sources[0].score == 80
    assert sources[0].text == "Full article text."
    assert sources[0].error is None

    # No URL at all for this one — build_sources_from_metacritic must not call fetch_full_text
    # (there's nothing to fetch), and must leave a clear, specific error rather than crashing.
    assert sources[1].outlet == "Outlet B"
    assert sources[1].text is None
    assert "did not provide a review URL" in sources[1].error

    assert progress_calls == [(1, 2, "Outlet A"), (2, 2, "Outlet B")]


def test_fetch_game_platforms_parses_platform_list(monkeypatch):
    fake_json = {
        "data": {
            "item": {
                "platforms": [
                    {"name": "PlayStation 5", "slug": "playstation-5",
                     "criticScoreSummary": {"reviewCount": 45}},
                    {"name": "PC", "slug": "pc", "criticScoreSummary": {"reviewCount": 12}},
                    # A platform release with zero critic reviews — should still be returned by
                    # fetch_game_platforms (it's the caller's job to filter these out), just
                    # with review_count 0 rather than crashing on a missing summary.
                    {"name": "Switch", "slug": "switch", "criticScoreSummary": {}},
                ]
            }
        }
    }

    class _FakeSession:
        def __init__(self):
            self.headers = {}

        def get(self, url, params=None, timeout=None):
            return _FakeReviewListResponse(fake_json)

    monkeypatch.setattr(metacritic.requests, "Session", _FakeSession)

    platforms = fetch_game_platforms("some-slug")

    assert platforms == [
        {"name": "PlayStation 5", "slug": "playstation-5", "review_count": 45},
        {"name": "PC", "slug": "pc", "review_count": 12},
        {"name": "Switch", "slug": "switch", "review_count": 0},
    ]


def test_fetch_review_list_all_platforms_dedupes_and_labels_duplicate_outlets(monkeypatch):
    # PS5 and PC each have reviews; Switch has zero critic reviews and must be skipped entirely
    # (no fetch_review_list call for it at all — real title behavior, not every platform release
    # gets reviewed).
    monkeypatch.setattr(metacritic, "fetch_game_platforms", lambda slug: [
        {"name": "PlayStation 5", "slug": "playstation-5", "review_count": 2},
        {"name": "PC", "slug": "pc", "review_count": 2},
        {"name": "Switch", "slug": "switch", "review_count": 0},
    ])

    def _fake_fetch_review_list(slug, platform=None):
        if platform == "playstation-5":
            return [
                # Outlet A reviewed the PS5 version and PC version separately — two distinct
                # URLs, both real reviews, both must be kept and disambiguated by platform.
                {"url": "https://a.example/ps5-review", "publicationName": "Outlet A",
                 "score": 90},
                # Outlet C's review URL covers both platforms at once (same URL will also come
                # back under the PC platform call below) — must be deduped down to one row.
                {"url": "https://c.example/multi-platform-review", "publicationName": "Outlet C",
                 "score": 75},
            ]
        if platform == "pc":
            return [
                {"url": "https://a.example/pc-review", "publicationName": "Outlet A",
                 "score": 85},
                {"url": "https://c.example/multi-platform-review", "publicationName": "Outlet C",
                 "score": 75},
                # Outlet B only ever reviewed the PC version — single platform, label untouched.
                {"url": "https://b.example/review", "publicationName": "Outlet B", "score": 70},
            ]
        raise AssertionError(f"fetch_review_list called for unexpected platform: {platform}")

    monkeypatch.setattr(metacritic, "fetch_review_list", _fake_fetch_review_list)

    items = fetch_review_list_all_platforms("some-slug")

    by_url = {i["url"]: i for i in items}
    assert len(items) == 4  # 5 raw items in, 1 dropped as a cross-platform URL duplicate

    # Outlet A: two genuinely separate reviews, each disambiguated by its own platform.
    assert by_url["https://a.example/ps5-review"]["publicationName"] == "Outlet A (PlayStation 5)"
    assert by_url["https://a.example/pc-review"]["publicationName"] == "Outlet A (PC)"

    # Outlet B: only ever appeared once post-dedupe — label left untouched, no platform suffix.
    assert by_url["https://b.example/review"]["publicationName"] == "Outlet B"

    # Outlet C: same URL appeared under both platform calls — deduped to a single row, and since
    # it's a single row post-dedupe, its label is also left untouched (not disambiguated).
    assert by_url["https://c.example/multi-platform-review"]["publicationName"] == "Outlet C"


def test_build_sources_from_metacritic_all_platforms_uses_consolidated_list(monkeypatch):
    fake_items = [
        {"url": "https://a.example/ps5-review", "publicationName": "Outlet A (PlayStation 5)",
         "score": 90, "date": "2024-01-01"},
        {"url": "https://a.example/pc-review", "publicationName": "Outlet A (PC)",
         "score": 85, "date": "2024-01-05"},
    ]
    monkeypatch.setattr(metacritic, "fetch_review_list_all_platforms", lambda slug: fake_items)
    monkeypatch.setattr(metacritic, "fetch_full_text",
                         lambda url, timeout=20: ("Full article text.", "requests", None))

    progress_calls = []
    sources = build_sources_from_metacritic_all_platforms(
        "some-slug",
        progress_cb=lambda done, total, outlet: progress_calls.append((done, total, outlet)),
    )

    assert len(sources) == 2
    assert {s.outlet for s in sources} == {"Outlet A (PlayStation 5)", "Outlet A (PC)"}
    assert all(s.text == "Full article text." and s.error is None for s in sources)
    assert progress_calls == [
        (1, 2, "Outlet A (PlayStation 5)"), (2, 2, "Outlet A (PC)"),
    ]


if __name__ == "__main__":
    import pytest
    raise SystemExit(pytest.main([__file__, "-v"]))
